"""
Convert placement report markdown files to a single formatted DOCX document.
Uses professional formatting with Calibri font and proper heading styles.
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT

# Configuration
REPORT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_FILE = os.path.join(REPORT_DIR, "Placement_Report_Laurentiu_Nae.docx")

# File order
CHAPTERS = [
    "README.md",
    "chapters/01-project-overview.md",
    "chapters/02-technical-architecture.md",
    "chapters/03-development-logbook.md",
    "chapters/04-features-implemented.md",
    "chapters/05-challenges-and-solutions.md",
    "chapters/06-technologies-and-tools.md",
    "chapters/07-testing-and-quality.md",
    "chapters/08-conclusions-and-future-work.md",
]


def setup_styles(doc):
    """Configure document styles with professional formatting."""

    # Set default font for Normal style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0, 0, 0)

    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(6)
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # Heading 1 - Chapter titles
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.page_break_before = True

    # Heading 2
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)

    # Heading 3
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Calibri'
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)

    # Heading 4
    h4 = doc.styles['Heading 4']
    h4.font.name = 'Calibri'
    h4.font.size = Pt(11)
    h4.font.bold = True
    h4.font.color.rgb = RGBColor(0, 0, 0)
    h4.paragraph_format.space_before = Pt(10)
    h4.paragraph_format.space_after = Pt(4)

    return doc


def add_title_page(doc):
    """Create a professional title page."""
    for _ in range(4):
        doc.add_paragraph()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Follicle Labeller")
    run.font.name = 'Calibri'
    run.font.size = Pt(32)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Medical Image Annotation Tool")
    run.font.name = 'Calibri'
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()
    doc.add_paragraph()

    # Report type
    report_type = doc.add_paragraph()
    report_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = report_type.add_run("Master's Placement Project Report")
    run.font.name = 'Calibri'
    run.font.size = Pt(16)
    run.bold = True

    for _ in range(6):
        doc.add_paragraph()

    # Author info
    info_lines = [
        ("Student:", "Laurentiu Nae"),
        ("Organization:", "Blockchain Advisors"),
        ("Period:", "January 2026"),
        ("Version:", "2.0.2"),
    ]

    for label, value in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = p.add_run(label + " ")
        run1.font.name = 'Calibri'
        run1.font.size = Pt(12)
        run1.bold = True
        run2 = p.add_run(value)
        run2.font.name = 'Calibri'
        run2.font.size = Pt(12)

    for _ in range(4):
        doc.add_paragraph()

    # Date
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run("January 29, 2026")
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.italic = True

    doc.add_page_break()


def add_formatted_run(paragraph, text, bold=False, italic=False, code=False):
    """Add a formatted run to a paragraph."""
    run = paragraph.add_run(text)
    run.font.name = 'Consolas' if code else 'Calibri'
    run.font.size = Pt(10) if code else Pt(11)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bold
    run.italic = italic
    return run


def parse_inline_formatting(paragraph, text):
    """Parse inline markdown formatting and add to paragraph."""
    # Remove links but keep text: [text](url) -> text
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)

    # Pattern to match inline formatting
    # Order matters: check longer patterns first
    pattern = r'(\*\*\*[^*]+\*\*\*|\*\*[^*]+\*\*|__[^_]+__|_[^_]+_(?![a-zA-Z0-9])|(?<![a-zA-Z0-9])\*[^*]+\*(?![a-zA-Z0-9])|`[^`]+`)'

    pos = 0
    for match in re.finditer(pattern, text):
        # Add text before match
        if match.start() > pos:
            add_formatted_run(paragraph, text[pos:match.start()])

        matched_text = match.group(0)

        # Bold and italic: ***text***
        if matched_text.startswith('***') and matched_text.endswith('***'):
            inner = matched_text[3:-3]
            add_formatted_run(paragraph, inner, bold=True, italic=True)
        # Bold: **text** or __text__
        elif (matched_text.startswith('**') and matched_text.endswith('**')):
            inner = matched_text[2:-2]
            add_formatted_run(paragraph, inner, bold=True)
        elif (matched_text.startswith('__') and matched_text.endswith('__')):
            inner = matched_text[2:-2]
            add_formatted_run(paragraph, inner, bold=True)
        # Italic: *text* or _text_
        elif (matched_text.startswith('*') and matched_text.endswith('*') and not matched_text.startswith('**')):
            inner = matched_text[1:-1]
            add_formatted_run(paragraph, inner, italic=True)
        elif (matched_text.startswith('_') and matched_text.endswith('_') and not matched_text.startswith('__')):
            inner = matched_text[1:-1]
            add_formatted_run(paragraph, inner, italic=True)
        # Code: `text`
        elif matched_text.startswith('`') and matched_text.endswith('`'):
            inner = matched_text[1:-1]
            add_formatted_run(paragraph, inner, code=True)
        else:
            add_formatted_run(paragraph, matched_text)

        pos = match.end()

    # Add remaining text
    if pos < len(text):
        add_formatted_run(paragraph, text[pos:])


def clean_all_markdown(text):
    """Remove ALL markdown formatting and return plain text."""
    # Remove links but keep text
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    # Remove bold ***text***
    text = re.sub(r'\*\*\*([^*]+)\*\*\*', r'\1', text)
    # Remove bold **text**
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    # Remove bold __text__
    text = re.sub(r'__([^_]+)__', r'\1', text)
    # Remove italic *text* (but not ** or file paths like *.js)
    text = re.sub(r'(?<![*\w])\*([^*]+)\*(?![*\w])', r'\1', text)
    # Remove italic _text_
    text = re.sub(r'(?<![_\w])_([^_]+)_(?![_\w])', r'\1', text)
    # Remove code `text`
    text = re.sub(r'`([^`]+)`', r'\1', text)
    return text


def create_table(doc, rows):
    """Create a formatted table."""
    if not rows or len(rows) < 1:
        return

    num_cols = max(len(row) for row in rows)
    if num_cols == 0:
        return

    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = row.cells[j]
                # Clean markdown from cell text
                clean_text = clean_all_markdown(cell_text)
                cell.text = clean_text
                # Format cell
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        if i == 0:  # Header row
                            run.bold = True

    doc.add_paragraph()


def process_markdown_file(doc, filepath):
    """Process a markdown file and add content to document."""
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    code_lines = []
    table_rows = []
    in_table = False

    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n\r')

        # Handle code blocks
        if line.startswith('```'):
            if in_code_block:
                # End code block
                if code_lines:
                    for code_line in code_lines:
                        p = doc.add_paragraph()
                        run = p.add_run(code_line)
                        run.font.name = 'Consolas'
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        p.paragraph_format.left_indent = Inches(0.25)
                        p.paragraph_format.space_after = Pt(0)
                    doc.add_paragraph()  # Space after code block
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Handle tables
        if '|' in line and line.strip().startswith('|'):
            cells = [c.strip() for c in line.split('|')]
            cells = [c for c in cells if c]

            # Check if it's a separator line
            if cells and all(re.match(r'^[-:]+$', c.strip()) for c in cells):
                i += 1
                continue

            if cells:
                table_rows.append(cells)
                in_table = True
            i += 1
            continue
        elif in_table:
            # End of table
            if table_rows:
                create_table(doc, table_rows)
            table_rows = []
            in_table = False
            # Don't increment i, process this line normally

        # Handle headings
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            text = clean_all_markdown(text)
            if level <= 4:
                doc.add_heading(text, level=level)
            else:
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
            i += 1
            continue

        # Handle horizontal rules
        if line.strip() in ['---', '***', '___']:
            p = doc.add_paragraph()
            run = p.add_run('─' * 50)
            run.font.color.rgb = RGBColor(150, 150, 150)
            i += 1
            continue

        # Handle bullet points (including nested)
        bullet_match = re.match(r'^(\s*)([-*+])\s+(.+)$', line)
        if bullet_match:
            indent = len(bullet_match.group(1))
            text = bullet_match.group(3)
            p = doc.add_paragraph(style='List Bullet')
            if indent >= 2:
                p.paragraph_format.left_indent = Inches(0.5)
            parse_inline_formatting(p, text)
            i += 1
            continue

        # Handle numbered lists
        num_match = re.match(r'^(\s*)(\d+)\.\s+(.+)$', line)
        if num_match:
            indent = len(num_match.group(1))
            text = num_match.group(3)
            p = doc.add_paragraph(style='List Number')
            if indent >= 2:
                p.paragraph_format.left_indent = Inches(0.5)
            parse_inline_formatting(p, text)
            i += 1
            continue

        # Handle regular paragraphs
        if line.strip():
            p = doc.add_paragraph()
            parse_inline_formatting(p, line)

        i += 1

    # Handle any remaining table
    if table_rows:
        create_table(doc, table_rows)


def main():
    """Main function to create the DOCX document."""
    print("Creating DOCX document...")

    doc = Document()
    setup_styles(doc)

    # Set page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    add_title_page(doc)

    for chapter_file in CHAPTERS:
        filepath = os.path.join(REPORT_DIR, chapter_file)
        if os.path.exists(filepath):
            print(f"Processing: {chapter_file}")
            process_markdown_file(doc, filepath)
        else:
            print(f"Warning: File not found: {filepath}")

    doc.save(OUTPUT_FILE)
    print(f"\nDocument saved to: {OUTPUT_FILE}")


if __name__ == "__main__":
    main()
